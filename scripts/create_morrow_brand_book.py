from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "MORROW_CHINA_Brand_Book.docx"
INK, JADE, ORANGE, CLOUD, CONCRETE, BLUE, WHITE, MUTED = "202522", "2E7568", "E26D45", "F7F7F2", "D9D8CF", "6C8F98", "FFFFFF", "65706B"

def rgb(h): return RGBColor.from_string(h)
def set_font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "STHeiti"; run.font.size = Pt(size); run.font.color.rgb = rgb(color); run.bold = bold; run.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "STHeiti")
def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), color); tcPr.append(shd)
def margins(cell):
    tcPr = cell._tc.get_or_add_tcPr(); mar = OxmlElement("w:tcMar")
    for side, val in (("top",100),("start",130),("bottom",100),("end",130)):
        x=OxmlElement("w:"+side); x.set(qn("w:w"),str(val)); x.set(qn("w:type"),"dxa"); mar.append(x)
    tcPr.append(mar)
def text_cell(cell, text, color=INK, bold=False, size=9.5):
    cell.text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.15; r=p.add_run(text); set_font(r,size,color,bold); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
def para(doc,text="",size=10.5,color=INK,bold=False,italic=False,align=None,before=0,after=6,line=1.25):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=line
    if align is not None:p.alignment=align
    if text:
        r=p.add_run(text); set_font(r,size,color,bold,italic)
    return p
def heading(doc,text,level=1):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20 if level==1 else 13); p.paragraph_format.space_after=Pt(7 if level==1 else 5)
    r=p.add_run(text); set_font(r,22 if level==1 else 14 if level==2 else 11.5, JADE if level<3 else ORANGE, level==3); return p
def bullet(doc,text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.18; r=p.add_run(text); set_font(r); return p
def table(doc,headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers): t.rows[0].cells[i].width=Inches(widths[i]); shade(t.rows[0].cells[i],JADE); text_cell(t.rows[0].cells[i],h,WHITE,True)
    for n,row in enumerate(rows):
        cs=t.add_row().cells
        for i,v in enumerate(row): cs[i].width=Inches(widths[i]); shade(cs[i],CLOUD if n%2==0 else WHITE); text_cell(cs[i],v)
    doc.add_paragraph().paragraph_format.space_after=Pt(2); return t
def main():
    doc=Document(); s=doc.sections[0]; s.top_margin=Inches(.7); s.bottom_margin=Inches(.65); s.left_margin=Inches(.85); s.right_margin=Inches(.85)
    normal=doc.styles["Normal"]; normal.font.name="STHeiti"; normal.font.size=Pt(10.5); normal._element.rPr.rFonts.set(qn("w:eastAsia"),"STHeiti")
    footer=s.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_font(footer.add_run("MORROW CHINA  /  BRAND BOOK  ·  2026"),8,MUTED,True)
    para(doc,"MORROW CHINA",11,JADE,True,align=WD_ALIGN_PARAGRAPH.CENTER,before=42,after=10)
    para(doc,"BRAND BOOK  /  01",10,ORANGE,True,align=WD_ALIGN_PARAGRAPH.CENTER,after=22)
    para(doc,"A new way into China.",34,JADE,False,align=WD_ALIGN_PARAGRAPH.CENTER,before=20,after=8,line=.95)
    para(doc,"面向当代人的中国入境游品牌",15,INK,False,align=WD_ALIGN_PARAGRAPH.CENTER,after=18)
    para(doc,"China, as it is becoming.",13,MUTED,False,True,WD_ALIGN_PARAGRAPH.CENTER,after=38)
    table(doc,["ROLE","CORE IDEA","VISUAL WORLD"],[["Private inbound travel","China in motion","Cloud White · Digital Jade · Signal Orange"]],[1.75,2.35,2.3]); doc.add_page_break()

    heading(doc,"01 这是什么品牌")
    para(doc,"MORROW CHINA 不是把景点剪得更热闹，而是把客人带到中国正在发生的地方：城市、食物、设计、山水、社区与日常生活。",12,INK,False,False,None,0,12,1.4)
    table(doc,["品牌关键词","观众应该感到"],[["Current / 当下","这里不是旧照片里的中国，而是今天仍在变化的中国。"],["Open / 开放","不猎奇、不端着，欢迎不同背景的人进入。"],["Human / 有人","旅行通过具体的人、餐桌、街道和关系发生。"],["Curious / 好奇","每个镜头都提出一个值得继续看的问题。"]],[1.9,4.45])
    heading(doc,"一句话定位",2); para(doc,"MORROW CHINA creates private journeys for travellers who want to meet China as it is becoming.",12,JADE,True,False,None,0,10,1.3)
    heading(doc,"与 AVIORA 的区别",2)
    table(doc,["AVIORA","MORROW CHINA"],[["安静、精致、有取舍","当代、开放、有探索感"],["像一本精装旅行刊物","像一部正在更新的中国城市纪录片"],["Considered China","China in motion"]],[3.2,3.2]); doc.add_page_break()

    heading(doc,"02 品牌战略")
    heading(doc,"品牌洞察",2); para(doc,"很多国际客人对中国的想象停留在古迹、熊猫和长城。他们真正想知道的是：今天的人如何生活？城市如何变化？传统怎样继续？中国的设计、食物、社区和年轻人正在发生什么？",11,INK,False,False,None,0,10,1.35)
    heading(doc,"使命",2); para(doc,"把中国正在发生的变化，转化成客人可以进入、理解并亲身参与的旅行体验。",12,JADE,True)
    heading(doc,"价值观",2)
    for x in ["Stay curious：保持好奇，不用旧答案解释新中国。","Be specific：用街区、店铺、人物、时间和真实选择说话。","Make room：给客人和当地人留下自然相遇的空间。","Move with change：接受城市、路线和生活持续变化。","Respect the everyday：日常生活不是背景，而是旅行的主体。"]: bullet(doc,x)
    heading(doc,"品牌本质",2); para(doc,"China in motion.\n正在移动、变化、生活中的中国。",18,JADE,False,True,None,3,8,1.05); doc.add_page_break()

    heading(doc,"03 语言系统")
    table(doc,["场景","推荐表达"],[["Tagline","A new way into China."],["辅助主张","China, as it is becoming."],["首页 H1","Meet China in motion."],["首页副文案","Private journeys through the cities, flavours, ideas and everyday lives shaping China now."],["CTA","Start with a conversation"],["CTA 说明","Tell us what you are curious about. We will suggest a route, a rhythm and a way in."]],[1.55,4.8])
    heading(doc,"写作原则",2)
    for x in ["从“看景点”改成“进入一个地方”。","从“完美旅程”改成“适合你的进入方式”。","从“隐藏宝藏”改成“一个仍在生活中的地方”。","从“体验当地文化”改成“和一个人、一顿饭、一条街发生关系”。"]: bullet(doc,x)
    heading(doc,"示例",2); para(doc,"不推荐：Discover the hidden gems of China in luxury.\n推荐：Spend a morning where the neighbourhood still belongs to its residents.",11,JADE,False,True,None,0,8,1.3); doc.add_page_break()

    heading(doc,"04 视觉识别")
    para(doc,"视觉概念：Tomorrow, in Plain Sight。不是未来科技，也不是复古中国，而是捕捉“未来已经出现在日常里”的瞬间。",11,INK,False,False,None,0,10,1.35)
    table(doc,["色名","HEX","用途"],[["Cloud White","#F7F7F2","主背景、字幕留白"],["Carbon","#202522","正文、深色底、主字幕"],["Digital Jade","#2E7568","品牌识别、按钮、地图线"],["Signal Orange","#E26D45","行动提示、时间、重点词"],["River Blue","#6C8F98","水、天空、辅助色"],["Warm Concrete","#D9D8CF","信息卡、分区底色"]],[1.6,1.4,3.3])
    para(doc,"色彩比例：Cloud White 45% / Carbon 25% / Digital Jade 15% / Warm Concrete 10% / Signal Orange 5%。Signal Orange 只能像城市里醒目的招牌一样点到为止。",10.5,MUTED,False,True,None,5,8,1.25)
    heading(doc,"字体",2)
    table(doc,["层级","英文字体","中文字体","用途"],[["Display","Space Grotesk","Noto Sans SC","首页标题、片头"],["Editorial","DM Sans","Noto Sans SC","正文、目的地故事"],["Mono","IBM Plex Mono","Noto Sans SC","日期、编号、城市代码"]],[1.2,1.7,1.7,1.7]); doc.add_page_break()

    heading(doc,"05 摄影与视频")
    heading(doc,"应该拍什么",2)
    para(doc,"城市更新、街边早餐、独立书店、茶馆、设计酒店、夜间交通、年轻创作者、菜市场、工艺工作室、社区公园、铁路和人与人的交谈。",11,INK,False,False,None,0,10,1.35)
    table(doc,["原则","执行"],[["真实胜过宏大","人物、动作、现场声优先于空旷地标。"],["细节胜过重复","门、手、食物、材质、交通和光线建立记忆。"],["变化胜过旧印象","展示新店、旧街、年轻人和传统的新用法。"],["尊重胜过猎奇","不把当地人当表演者；拍摄前取得同意。"]],[1.9,4.45])
    heading(doc,"视频结构",2)
    table(doc,["时长","结构","节奏"],[["15 秒","当代瞬间 → 具体观察 → 品牌落版","平均 1.5–2.5 秒/镜头"],["30 秒","进入城市 → 遇见人物 → 一项体验 → CTA","平均 2–3.5 秒/镜头"],["60–90 秒","一个问题 → 三个现场章节 → 新的理解","人物/环境声镜头 4–6 秒"]],[.9,3.5,1.95]); doc.add_page_break()

    heading(doc,"06 声音、字幕与动态")
    heading(doc,"音乐",2); para(doc,"选择轻快但不喧闹的当代音乐：电子纹理、木吉他、钢琴、轻打击和真实城市声。不要使用“异国情调”的廉价民族配乐。",11,INK,False,False,None,0,9,1.35)
    heading(doc,"环境声",2); para(doc,"扫码声、地铁进站、锅铲、脚步、街头叫卖、风和对话片段都应被保留。MORROW CHINA 的真实感来自声音，而不只是画面。",11,JADE,True,False,None,0,9,1.3)
    heading(doc,"字幕",2)
    for x in ["英文每行 32–42 个字符；中文每行 12–16 个字；最多两行。","字幕使用 Carbon 或 Cloud White；不要描边发光。","关键词可用 Digital Jade 或 Signal Orange，一条字幕最多强调一个词。","出现/消失 180–260ms 淡入淡出；不使用打字机和弹跳效果。"]: bullet(doc,x)
    heading(doc,"推荐字幕",2); para(doc,"The future is already here.\n未来，已经在日常里发生。",15,JADE,False,True,None,0,7,1.1); para(doc,"Not a checklist. A way in.\n不是景点清单，而是一种进入方式。",15,JADE,False,True,None,0,7,1.1); doc.add_page_break()

    heading(doc,"07 产品与内容落地")
    heading(doc,"产品命名",2)
    for x in ["A Morning in Shanghai’s Old Neighbourhoods","Chengdu After Dark: Tea, Spice and Street Life","Beijing, Beyond the Monument","Three Days in China’s Design Cities"]: bullet(doc,x)
    heading(doc,"服务原则",2)
    for x in ["每条路线至少包含一个真实生活场景、一个文化解释场景和一个自由探索段落。","给客人选择进入深度的方式：看、听、吃、走、聊。","价格、酒店、交通、向导和取消规则写清楚；当代不等于随意。","共享法律主体和执行团队时，在报价与合同中清楚说明承接主体。"]: bullet(doc,x)
    heading(doc,"内容支柱",2)
    table(doc,["支柱","内容方向"],[["Now in China","城市、设计、食物和生活变化"],["Ways In","清晨、夜晚、步行、餐桌、铁路"],["People of Place","向导、厨师、设计师、店主、工匠"],["Practical China","支付、铁路、门票、登记、网络"]],[1.75,4.6]); doc.add_page_break()

    heading(doc,"08 启动与注册提醒")
    heading(doc,"30 天启动",2)
    table(doc,["时间","动作"],[["第 1 周","确认名称、商标检索、域名、Logo、字体和色板。"],["第 2 周","整理 3 个城市的当代生活素材，建立视频模板。"],["第 3 周","制作首页、3 条旗舰产品、6 条短视频和 2 篇 Now in China 内容。"],["第 4 周","测试咨询表单、CTA、完播率和双品牌受众区隔。"]],[1.25,5.1])
    heading(doc,"注册前提示",2); para(doc,"MORROW CHINA 目前是品牌策略候选名，不代表已获得商标权。正式使用前，应至少核查中国、英国、美国和欧盟的文字商标与近似商标，重点关注第 39 类，并单独检索 MORROW。",11,ORANGE,True,False,None,0,12,1.35)
    para(doc,"MORROW CHINA\nA new way into China.",20,JADE,True,True,WD_ALIGN_PARAGRAPH.CENTER,55,6,1.0)
    para(doc,"China, as it is becoming.",12,MUTED,False,True,WD_ALIGN_PARAGRAPH.CENTER)
    OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
if __name__=='__main__': main()
