from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "AVIORA_Video_Visual_Guide.docx"
LOGO = ROOT / "public" / "brand" / "aviora-logo.jpg"

INK = "1B1C19"
PINE = "30483D"
COPPER = "A66A4B"
RICE = "F5F1E8"
MIST = "DDE4DC"
MUTED = "6F6C66"
WHITE = "FFFFFF"

def rgb(h):
    return RGBColor.from_string(h)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color)

def cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_cell_text(cell, text, color=INK, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Arial Unicode MS"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(color)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)

def set_run(run, font="Arial Unicode MS", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Arial Unicode MS"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic

def para(doc, text="", style=None, size=10.5, color=INK, bold=False, italic=False, align=None, before=0, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p

def heading(doc, text, level=1):
    sizes = {1: 22, 2: 15, 3: 11.5}
    colors = {1: PINE, 2: PINE, 3: COPPER}
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt({1: 22, 2: 15, 3: 10}[level])
    p.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}[level])
    r = p.add_run(text)
    set_run(r, font="Cormorant Garamond" if level < 3 else "Manrope", size=sizes[level], color=colors[level], bold=level == 3)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run(r)
    return p

def page_break(doc):
    doc.add_page_break()

def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        shade(c, PINE)
        set_cell_text(c, h, color=WHITE, bold=True, size=9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            shade(cells[i], RICE if len(t.rows) % 2 == 0 else "FFFFFF")
            set_cell_text(cells[i], value, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial Unicode MS"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    # Footer
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("AVIORA  /  VIDEO VISUAL GUIDE  ·  2026")
    set_run(fr, size=8, color=MUTED, bold=True)

    # Cover
    p = para(doc, "AVIORA", size=11, color=PINE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=38, after=10)
    if LOGO.exists():
        doc.add_picture(str(LOGO), width=Inches(2.25))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "VIDEO  /  VISUAL GUIDE", size=10, color=COPPER, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=22)
    para(doc, "Quiet Cartography", size=34, color=PINE, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=6, line=0.95)
    para(doc, "给视频剪辑师的品牌执行手册", size=15, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    para(doc, "让中国被看见，也被理解。", size=13, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=38)
    table(doc, ["用途", "版本", "核心色彩"], [["品牌视频 / 目的地片 / 社媒短片", "V1.0 · 2026-08-06", "Rice Paper · Pine Shadow · Oxidised Copper"]], [2.35, 1.55, 2.5])
    page_break(doc)

    heading(doc, "01 先记住 AVIORA 是什么", 1)
    para(doc, "AVIORA 不是把景点剪得更热闹，而是让观众感到：中国很大、很深，但有人知道如何带你进入其中。视频的任务不是证明我们有多少景点，而是证明我们有判断、有分寸、有人照应。", size=12, color=INK, after=12, line=1.4)
    table(doc, ["关键词", "画面上要让人感到"], [
        ["Considered / 被认真想过", "镜头有选择，剪辑不贪多，留出观察时间。"],
        ["Grounded / 在地", "真实街道、真实交通、真实人物和具体细节。"],
        ["Cultured / 有文化", "有语境的声音与文字，不靠符号堆砌。"],
        ["Cared for / 被照应", "有人迎接、解释、等待、确认和处理变化。"],
    ], [1.9, 4.5])
    heading(doc, "一句话判断", 2)
    para(doc, "如果一个镜头只是“漂亮”，却没有地点、关系、时间或情绪，它还不够 AVIORA。", size=12, color=PINE, bold=True, after=10)
    para(doc, "推荐英文品牌句：Thoughtful private journeys across China, designed and delivered by people who know it from the inside.", size=10.5, color=MUTED, italic=True)
    page_break(doc)

    heading(doc, "02 色彩：像纸、墨、松影和氧化铜", 1)
    para(doc, "色彩要低饱和、耐看、有材质感。不要把“高级”理解成黑金；AVIORA 的高级来自克制和留白。", after=10)
    table(doc, ["色名", "HEX", "剪辑使用"], [
        ["Rice Paper", "#F5F1E8", "字幕底板、片头留白、行程/目的地信息页"],
        ["Ink Black", "#1B1C19", "正文字幕、导航、深色片尾、主文字"],
        ["Pine Shadow", "#30483D", "品牌识别色、章节标题、地图线、CTA"],
        ["Oxidised Copper", "#A66A4B", "编号、时间、重点词、极少量强调"],
        ["Celadon Mist", "#DDE4DC", "信息卡、信任内容、柔和过渡背景"],
    ], [1.55, 1.25, 3.6])
    heading(doc, "调色标准", 2)
    bullet(doc, "整体饱和度下调 10–20%，保留真实肤色和食物颜色；不要套统一橙色滤镜。")
    bullet(doc, "黑位不要压死；保留建筑、树影和室内暗部的层次。")
    bullet(doc, "高光偏暖但不发黄；绿色偏松柏，不要荧光青。")
    bullet(doc, "Copper 只作为视觉标点，画面占比建议低于 8%。")
    heading(doc, "禁止", 2)
    para(doc, "金色渐变、红黑中国风、霓虹旅游色、过度 HDR、过饱和青橙、明显胶片边框和廉价旅行贴纸。", color=COPPER, bold=True)
    page_break(doc)

    heading(doc, "03 字体与字幕", 1)
    para(doc, "字体是 AVIORA 最稳定的识别元素。标题有文学性，信息有秩序；两者必须同时存在。", after=10)
    table(doc, ["层级", "英文字体", "中文字体", "参数"], [
        ["片头 / 大标题", "Cormorant Garamond", "Noto Serif SC", "Regular / 400–500；行高 0.95"],
        ["目的地叙事", "Newsreader", "Noto Serif SC", "Regular；不要全大写"],
        ["字幕 / 信息", "Manrope", "Noto Sans SC", "400–600；行高 1.35–1.5"],
        ["日期 / 编号", "IBM Plex Mono", "Noto Sans SC", "小面积；用于 01、09 DAYS、SHANGHAI"],
    ], [1.45, 1.7, 1.55, 1.7])
    heading(doc, "字幕规范", 2)
    bullet(doc, "每行英文建议 32–42 个字符；中文每行 12–16 个字。最多两行。")
    bullet(doc, "字幕不要贴底；短视频底部安全区至少 8% 画面高度。")
    bullet(doc, "主字幕用 Rice Paper 文字或 Ink Black 半透明底板；不要描边发光。")
    bullet(doc, "关键词可用 Pine Shadow 或 Copper，一条字幕最多强调一个词。")
    bullet(doc, "字幕出现和消失用 180–260ms 淡入淡出；不使用打字机、弹跳和霓虹效果。")
    heading(doc, "推荐字幕句式", 2)
    para(doc, "Not more places. Better chosen places.\n不是去更多地方，而是把地方选得更对。", size=13, color=PINE, italic=True, after=4, line=1.15)
    para(doc, "The city changes when you slow down.\n当你慢下来，城市才开始显出层次。", size=13, color=PINE, italic=True, line=1.15)
    page_break(doc)

    heading(doc, "04 镜头语言：少一点证明，多一点观察", 1)
    table(doc, ["优先拍什么", "怎么拍", "避免什么"], [
        ["人与人的关系", "向导解释、客人倾听、主人递茶、司机等待；中近景，保留动作。", "正面摆拍、对镜头挥手、统一笑脸。"],
        ["时间感", "清晨开门、车窗移动、午后光线、夜晚收束；允许镜头停留。", "景点快闪、每秒一个地标。"],
        ["材质与细节", "门轴、石阶、茶汤、车票、手写字、酒店窗帘；微距与环境声。", "无意义的装饰特写、素材库式 B-roll。"],
        ["空间关系", "从街巷到建筑，从城市到山水；用横移、推拉或固定镜头建立尺度。", "过度无人机、无地点语境的风景拼贴。"],
    ], [1.45, 3.2, 1.75])
    heading(doc, "镜头节奏", 2)
    para(doc, "默认剪辑速度：平均单镜头 2.5–4.5 秒；关键观察镜头 5–8 秒；转场镜头 1.5–2 秒。不要把每条片都剪成广告预告片。", size=11.5, color=INK, bold=True, after=10)
    bullet(doc, "开场 0–3 秒：一个有地点感的动作，不要先放 logo。")
    bullet(doc, "3–12 秒：建立人物/路线/时间中的一个关系。")
    bullet(doc, "中段：用 2–3 个章节推进，不超过 5 个视觉主题。")
    bullet(doc, "结尾：回到人或空间，再出现 AVIORA 和 CTA。")
    page_break(doc)

    heading(doc, "05 音乐、声音与转场", 1)
    heading(doc, "音乐方向", 2)
    bullet(doc, "当代、留白、轻微东方质感；优先钢琴、木管、弦乐、环境氛围和极简打击。")
    bullet(doc, "避免旅游宣传片式大合唱、民族符号堆叠、过强鼓点和情绪强行上扬。")
    bullet(doc, "音乐不应替代画面解释；让环境声承担真实感。")
    heading(doc, "必须保留的环境声", 2)
    para(doc, "脚步、车门、站台提示音、茶水、街道远声、风、鸟、筷子与餐具、门开合、向导的一句原声。环境声建议比音乐高 1–2 dB 的存在感。", after=10)
    heading(doc, "转场", 2)
    table(doc, ["推荐", "慎用"], [
        ["动作匹配、视线匹配、自然遮挡、声音先行、固定镜头切换、短暂留白", "大面积光效、旋转、故障、速度拉伸、旅游模板转场"],
    ], [3.2, 3.2])
    heading(doc, "片头与片尾", 2)
    bullet(doc, "片头：先进入中国，再进入品牌。Logo 最早在 3 秒后出现。")
    bullet(doc, "片尾：Rice Paper 或深 Pine Shadow 背景；AVIORA 字标 + 一句主张 + 一个明确 CTA。")
    bullet(doc, "片尾不要堆电话、邮箱、社媒和长免责声明；信息分两层呈现。")
    page_break(doc)

    heading(doc, "06 视频结构模板", 1)
    table(doc, ["时长", "结构", "画面任务", "文字任务"], [
        ["15 秒", "观察 → 细节 → 品牌", "1 个地点动作 + 2 个细节 + 1 个收束", "最多 2 句字幕 + logo"],
        ["30 秒", "进入 → 理解 → 选择 → 行动", "人物关系、空间变化、一个服务证据", "3–4 个章节短句 + CTA"],
        ["60–90 秒", "命题 → 路线 → 人 → 体验 → 回望", "建立完整情绪弧线，保留环境声", "旁白可用；每段只说一个判断"],
        ["目的地片", "地点性 → 生活性 → AVIORA 方式", "不要只展示景点，要展示怎么进入它", "标题含地点；结尾含线路方向"],
    ], [0.85, 1.55, 2.7, 1.3])
    heading(doc, "旁白语气", 2)
    para(doc, "像一个真正到过这里的人在分享判断，而不是播报百科。句子短，留有停顿；不要每个镜头都解释。", after=10)
    para(doc, "示例：\nBeijing is not only a city of monuments.\nIt is also the quiet hour before the gates open.\n北京不只是宏大的古迹，也有城门开启前的那一小时。", size=12.5, color=PINE, italic=True, line=1.2)
    page_break(doc)

    heading(doc, "07 Logo、片尾与交付检查", 1)
    heading(doc, "Logo 使用", 2)
    bullet(doc, "优先使用透明背景 SVG/PNG；白底 JPG/WebP 只用于白色背景的静态封面。")
    bullet(doc, "Logo 周围保留至少一个字标高度的安全距离。")
    bullet(doc, "不拉伸、不倾斜、不加阴影、不加金色描边、不与地标贴合变形。")
    heading(doc, "片尾标准", 2)
    para(doc, "AVIORA\nThoughtful private journeys across China.\nShape your China journey →", size=18, color=PINE, bold=True, after=12, line=1.0)
    heading(doc, "交付前检查表", 2)
    for item in [
        "画面低饱和但不灰；肤色、食物和自然色真实。",
        "没有模板化旅游转场、炫技特效或无语境地标拼贴。",
        "字幕字体、行数、位置、颜色和安全区符合本手册。",
        "环境声可听见；音乐没有压过人物和地点。",
        "Logo 使用正确；片尾 CTA 清楚且不拥挤。",
        "导出 16:9 主版、9:16 竖版和无字幕 clean version；文件名包含日期、比例和版本。",
    ]:
        bullet(doc, item)
    heading(doc, "推荐文件命名", 2)
    para(doc, "AVIORA_[主题]_[比例]_[语言]_[版本]_[日期].mp4\n例：AVIORA_ChengduTea_9x16_EN_v03_20260806.mp4", size=10.5, color=MUTED, italic=True)
    page_break(doc)

    heading(doc, "08 参考：AVIORA 的一句话", 1)
    para(doc, "AVIORA 让复杂的中国变得可进入，让珍贵的时间被更好地使用。", size=22, color=PINE, italic=True, line=1.05, after=18)
    para(doc, "请把这句话当作剪辑判断，而不是必须出现的广告语：每个镜头、每个停顿、每个声音，都应该让观众更接近真实的中国，而不是更接近一条普通旅游广告。", size=12, color=INK, line=1.45)
    para(doc, "AVIORA  ·  China Prime DMC\nPrivate China travel, designed and delivered in China.", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=60, line=1.3)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__":
    main()
