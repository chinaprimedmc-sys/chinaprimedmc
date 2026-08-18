from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT=Path(__file__).resolve().parents[1]
OUT=ROOT/'docs'/'AVIORA_MORROW_Social_Profile_Pack.docx'
INK='202522'; JADE='2E7568'; ORANGE='E26D45'; CLOUD='F7F7F2'; WHITE='FFFFFF'; MUTED='65706B'; BLUE='6C8F98'
def rgb(h): return RGBColor.from_string(h)
def font(r,size=10.5,color=INK,bold=False,italic=False):
    r.font.name='STHeiti'; r.font.size=Pt(size); r.font.color.rgb=rgb(color); r.bold=bold; r.italic=italic; r._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),'STHeiti')
def shade(cell,color):
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); cell._tc.get_or_add_tcPr().append(shd)
def margins(cell):
    mar=OxmlElement('w:tcMar')
    for side,val in [('top',90),('start',130),('bottom',90),('end',130)]:
        x=OxmlElement('w:'+side); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); mar.append(x)
    cell._tc.get_or_add_tcPr().append(mar)
def para(doc,text='',size=10.5,color=INK,bold=False,italic=False,align=None,before=0,after=6,line=1.25):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=line
    if align is not None:p.alignment=align
    if text: font(p.add_run(text),size,color,bold,italic)
    return p
def h(doc,text,level=1):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20 if level==1 else 14); p.paragraph_format.space_after=Pt(7 if level==1 else 5); font(p.add_run(text),22 if level==1 else 14,JADE if level<3 else ORANGE,level==3); return p
def bullet(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.18; font(p.add_run(text)); return p
def field(doc,label,value,note=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2); font(p.add_run(label.upper()),9,ORANGE,True)
    p2=doc.add_paragraph(); p2.paragraph_format.space_after=Pt(3); p2.paragraph_format.line_spacing=1.25; font(p2.add_run(value),11,INK,False)
    if note: para(doc,'用途：'+note,9,MUTED,False,True,after=5)
def table(doc,headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,x in enumerate(headers):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]); shade(c,JADE); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; c.text=''; font(c.paragraphs[0].add_run(x),9,WHITE,True)
    for n,row in enumerate(rows):
        cs=t.add_row().cells
        for i,x in enumerate(row):
            c=cs[i]; c.width=Inches(widths[i]); shade(c,CLOUD if n%2==0 else WHITE); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; c.text=''; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.15; font(p.add_run(x),9.4)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def brand_intro(doc,name,role,color,desc):
    h(doc,name,1); para(doc,role,12,color,True,after=5); para(doc,desc,10.5,MUTED,False,False,None,0,9,1.3)
def add_brand(doc,name,role,color,facebook,instagram,whatsapp,cover,pinned):
    brand_intro(doc,name,role,color,'本页全部文字可直接复制到对应平台后台；括号内为内部操作说明，不要复制到公开页面。')
    h(doc,'Facebook Page',2)
    for label,value,note in facebook: field(doc,label,value,note)
    field(doc,'封面图片方向',cover,'给设计或视频团队的视觉说明。')
    field(doc,'置顶帖',pinned,'发布后固定在主页顶部。')
    h(doc,'Instagram Profile',2)
    for label,value,note in instagram: field(doc,label,value,note)
    h(doc,'WhatsApp Business',2)
    for label,value,note in whatsapp: field(doc,label,value,note)
    doc.add_page_break()
def main():
    d=Document(); s=d.sections[0]; s.top_margin=Inches(.7); s.bottom_margin=Inches(.65); s.left_margin=Inches(.85); s.right_margin=Inches(.85)
    normal=d.styles['Normal']; normal.font.name='STHeiti'; normal.font.size=Pt(10.5); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'STHeiti')
    footer=s.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(footer.add_run('SOCIAL PROFILE PACK  ·  2026'),8,MUTED,True)
    para(d,'AVIORA  +  MORROW CHINA',11,JADE,True,WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.CENTER,38,10)
    para(d,'SOCIAL MEDIA PROFILE PACK',10,ORANGE,True,False,WD_ALIGN_PARAGRAPH.CENTER,0,23)
    para(d,'平台名称、简介与自动回复',30,JADE,False,False,WD_ALIGN_PARAGRAPH.CENTER,10,8,.95)
    para(d,'Facebook  ·  Instagram  ·  WhatsApp Business',13,MUTED,False,True,WD_ALIGN_PARAGRAPH.CENTER,0,35)
    table(d,['品牌','定位','使用方法'],[['AVIORA','安静、精致、私人定制','复制对应字段到平台后台'],['MORROW CHINA','当代、开放、城市生活','不要与 AVIORA 混用文案']],[1.55,2.45,2.4])
    para(d,'版本 1.0  ·  2026-08-06',9,MUTED,False,True,WD_ALIGN_PARAGRAPH.CENTER,35,0)
    d.add_page_break()
    add_brand(d,'AVIORA','Thoughtful private journeys across China.',JADE,
      [('Page name','AVIORA | Private China Travel','Facebook 页面公开名称。若竖线不被允许，改用 AVIORA Private China Travel。'),('Username','@avioraprivatechina','优先使用；如果不可用，依次尝试 @aviorachina / @aviorajourneys。'),('Category','Travel Company','Facebook 页面分类。'),('Short bio','Thoughtful private journeys across China, designed and delivered by people who know it from the inside.','短简介；保持英文。'),('About','AVIORA creates private journeys across China for travellers who value good judgement, personal service and a more considered pace. From Beijing and Xi’an to Shanghai, Chengdu, Yunnan and beyond, we design routes around your interests, comfort and time. Carefully chosen hotels, private guides, local transport and China-based support are arranged by one accountable team. No forced shopping. No rushed checklist. Just a clearer, more personal way to experience China. AVIORA is the international travel brand of China Prime DMC, operated in China by a licensed inbound tourism operator.','Facebook 关于我们长简介。'),('CTA button','Send Message','主页主按钮。'),('Email','chinaprimedmc@gmail.com','公开联系邮箱。'),('WhatsApp link','https://wa.me/447985052302','按钮或联系信息中使用。')],
      [('Instagram name','AVIORA | Private China Travel','Instagram 显示名称。'),('Username','@avioraprivatechina','与 Facebook 尽量一致。'),('Bio','Thoughtful private journeys across China.\nSelected hotels · Private guides · China-based support\nNo forced shopping.','Instagram 150 字符内简介；换行可保留。'),('Link in bio','https://www.chinaprimedmc.com/contact','链接到咨询页。')],
      [('Business name','AVIORA Private China Travel','WhatsApp Business 显示名称。'),('Category','Travel Agency','业务分类。'),('Description','Private China journeys designed around your interests, pace and comfort. Local specialists, selected hotels, private guides and support throughout your trip.','Business description。'),('Greeting message','Hello, and welcome to AVIORA.\n\nThank you for reaching out about travelling in China. Please share your approximate dates, number of travellers, places you are considering and what matters most to you.\n\nA China specialist will review your ideas and suggest a clear first direction.','新客户首次发消息自动发送。'),('Away message','Thank you for your message. Our team is currently offline, but we have received your enquiry. We will reply as soon as possible during working hours. If your travel is already underway and you need urgent assistance, please write URGENT at the beginning of your message.','非工作时间自动回复。'),('/welcome','Welcome to AVIORA. We design private China journeys around your interests, preferred pace and comfort.','快捷回复。'),('/planning','We begin with your dates, travel style, interests and preferred rhythm. From there, we recommend a route rather than sending a standard package.','快捷回复。'),('/licensed','Your China services are arranged and delivered in China by a licensed inbound tourism operator. Contracting company details are provided clearly in your proposal and booking documents.','快捷回复。'),('/shopping','Our private journeys do not include compulsory shopping stops. Shopping can be added only when it is genuinely relevant to your interests.','快捷回复。'),('/next','Once we understand your dates and priorities, we will suggest a first route direction, recommended stay length and service approach.','快捷回复。')],
      '低饱和山水、建筑、清晨城市；Rice Paper、Pine Shadow、墨色；画面安静、有留白、有细节。',
      'China is not one destination. It is a collection of cities, landscapes, flavours and everyday lives — each with its own rhythm. AVIORA designs private journeys around the parts of China you are most curious about, with carefully chosen hotels, private guides and local support throughout. Tell us what you would like to see, how you like to travel and what matters most. We will suggest a thoughtful way in.')
    add_brand(d,'MORROW CHINA','A new way into China.',ORANGE,
      [('Page name','MORROW CHINA | Contemporary China Travel','Facebook 页面公开名称。若过长，改用 MORROW CHINA Travel。'),('Username','@morrowchina','优先使用；备选 @morrowchinatravel / @morrowchina.travel。'),('Category','Travel Company','Facebook 页面分类。'),('Short bio','A new way into China — through cities, food, design and everyday life.','短简介；保持英文。'),('About','MORROW CHINA creates private journeys for travellers who want to meet China as it is becoming. Go beyond the usual checklist through neighbourhoods, local food, independent spaces, design, contemporary culture and the people shaping everyday life. We connect the practical details — hotels, guides, transport and timing — with the moments that make a place feel real. This is China in motion: open, specific, human and always changing.','Facebook 关于我们长简介。'),('CTA button','Send Message','主页主按钮。'),('Email','chinaprimedmc@gmail.com','公开联系邮箱。'),('WhatsApp link','https://wa.me/447985052302','按钮或联系信息中使用。')],
      [('Instagram name','MORROW CHINA | Contemporary China Travel','Instagram 显示名称。'),('Username','@morrowchina','与 Facebook 尽量一致。'),('Bio','A new way into China.\nCities · food · design · everyday life\nPrivate journeys, made locally.','Instagram 150 字符内简介。'),('Link in bio','https://www.chinaprimedmc.com/contact','链接到咨询页。')],
      [('Business name','MORROW CHINA Travel','WhatsApp Business 显示名称。'),('Category','Travel Agency','业务分类。'),('Description','Private journeys through the cities, flavours, ideas and everyday lives shaping China now.','Business description。'),('Greeting message','Hi, and welcome to MORROW CHINA.\n\nWhat part of China are you curious about right now — city life, food, design, history, landscapes or everyday culture?\n\nTell us your approximate dates, who is travelling and what you would like to discover. We will suggest a route and a way in.','新客户首次发消息自动发送。'),('Away message','Thanks for writing to MORROW CHINA. We have received your message and will reply during working hours. If you are already travelling in China and need urgent help, please begin your message with URGENT.','非工作时间自动回复。'),('/welcome','Welcome to MORROW CHINA. We create private journeys through the China that is changing, living and evolving now.','快捷回复。'),('/style','Our routes often combine neighbourhoods, food, design, local culture and well-known landmarks. The balance depends on what you are most curious about.','快捷回复。'),('/planning','We do not begin with a fixed package. We begin with your interests, dates and preferred rhythm, then recommend a way into China.','快捷回复。'),('/practical','We arrange the practical side too — selected hotels, private guides, transport, tickets and local support throughout.','快捷回复。'),('/shopping','Our journeys do not include compulsory shopping stops. Shopping or market visits are included because they are relevant to your interests.','快捷回复。'),('/next','Send us your dates, the people travelling and three things you would like to understand or experience in China. We will reply with a first route direction.','快捷回复。')],
      '城市街道、早餐、夜间交通、设计空间、人物和生活细节；Cloud White、Digital Jade、Signal Orange；画面更近、更有现场感。',
      'China is changing quickly. New restaurants open beside old noodle shops. Historic neighbourhoods become creative districts. Traditional craft finds new forms. MORROW CHINA creates private journeys for travellers who want to see that China — not just the China on a checklist. Tell us what you are curious about. We will suggest a route, a rhythm and a way in.')
    h(d,'最后检查',1)
    for x in ['Facebook、Instagram 和 WhatsApp 的品牌名称保持一致。','两个品牌不要共用同一套头像、封面和简介。','公开页面只使用英文；中文说明保留给内部团队。','WhatsApp 账号上线前确认号码、营业时间和自动回复触发条件。','MORROW CHINA 正式使用前完成商标、域名和账号核查。']: bullet(d,x)
    d.save(OUT); print(OUT)
if __name__=='__main__': main()
